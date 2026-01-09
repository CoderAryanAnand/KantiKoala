from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

def create_final_subjonctif_docx():
    doc = Document()

    # Title
    doc.add_heading('Beispiel Lerninhalt: Französisch Subjonctif (Komplett)', 0)

    # 1. Structure/Formation
    doc.add_heading('1. Bildung (Formation)', level=1)
    doc.add_paragraph('Der Subjonctif wird vom Stamm der 3. Person Plural (ils/elles) des Präsens abgeleitet.')

    # 1.1 Step-by-step
    doc.add_heading('1.1 Allgemeine Regel', level=2)
    doc.add_paragraph('Schritt-für-Schritt-Anleitung:')
    
    doc.add_paragraph('Nehmen Sie die "ils"-Form im Präsens (z.B. ils finissent).', style='List Number')
    doc.add_paragraph('Streichen Sie die Endung -ent (Stamm: finiss-).', style='List Number')
    doc.add_paragraph('Hängen Sie die Subjonctif-Endungen an:', style='List Number')

    # Endings list
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('je: -e, tu: -es, il: -e').bold = True
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('nous: -ions, vous: -iez').bold = True
    p.add_run(' (siehe Imparfait-Regel unten)').italic = True
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('ils: -ent').bold = True

    # 1.2 Regular Verbs Examples
    doc.add_heading('1.2 Beispiele regelmässiger Konjugation', level=2)
    
    table_reg = doc.add_table(rows=1, cols=4)
    table_reg.style = 'Table Grid'
    hdr_cells = table_reg.rows[0].cells
    hdr_cells[0].text = 'Gruppe'
    hdr_cells[1].text = 'Infinitiv'
    hdr_cells[2].text = 'Stamm'
    hdr_cells[3].text = 'Ergebnis (je / nous)'

    data_reg = [
        ('-er Verben', 'parler', 'parl-', 'que je parle\nque nous parlions'),
        ('-ir (Typ finir)', 'finir', 'finiss-', 'que je finisse\nque nous finissions'),
        ('-re Verben', 'attendre', 'attend-', "que j'attende\nque nous attendions"),
        ('-ir (Typ dormir)', 'dormir', 'dorm-', 'que je dorme\nque nous dormions'),
    ]

    for group, inf, stamm, res in data_reg:
        row = table_reg.add_row().cells
        row[0].text = group
        row[1].text = inf
        row[2].text = stamm
        row[3].text = res

    # 1.3 Imparfait Rule (NEW)
    doc.add_heading('1.3 Besonderheit: Die Imparfait-Regel', level=2)
    doc.add_paragraph('Für die Formen von "nous" und "vous" gibt es eine wichtige Merkhilfe:')
    
    p_imp = doc.add_paragraph(style='List Bullet')
    p_imp.add_run('Die Endungen ').text
    p_imp.add_run('-ions').bold = True
    p_imp.add_run(' und ').text
    p_imp.add_run('-iez').bold = True
    p_imp.add_run(' sind identisch mit dem Imparfait.').text
    
    doc.add_paragraph('Bei regelmässigen Verben gleicht der Subjonctif daher oft dem Imparfait:', style='List Bullet')
    doc.add_paragraph('   Imparfait: nous parlions  |  Subjonctif: que nous parlions (Identisch)', style='List Paragraph')
    
    p_warn = doc.add_paragraph(style='List Bullet')
    p_warn.add_run('Achtung: ').bold = True
    p_warn.add_run('Dies gilt nicht, wenn der Stamm unregelmässig ist (z.B. faire -> fassions, aber Imparfait = faisions).').text

    # 2. Irregular Tables
    doc.add_heading('2. Unregelmässige Verben', level=1)
    
    table_irr = doc.add_table(rows=1, cols=3)
    table_irr.style = 'Table Grid'
    hdr_irr = table_irr.rows[0].cells
    hdr_irr[0].text = 'Verb'
    hdr_irr[1].text = 'Stamm'
    hdr_irr[2].text = 'Konjugation'

    data_irr = [
        ('être', 'sois/soy', 'que je sois, que nous soyons'),
        ('avoir', 'ai/ay', "que j'aie, que nous ayons"),
        ('aller', 'aill/all', "que j'aille, que nous allions"),
        ('faire', 'fass-', 'que je fasse, que nous fassions'),
        ('pouvoir', 'puiss-', 'que je puisse, que nous puissions'),
        ('savoir', 'sach-', 'que je sache, que nous sachions')
    ]

    for verb, stamm, bsp in data_irr:
        row = table_irr.add_row().cells
        row[0].text = verb
        row[1].text = stamm
        row[2].text = bsp

    # 3. Usage
    doc.add_heading('3. Verwendung (L’emploi)', level=1)
    doc.add_paragraph('Der Subjonctif steht meist nach "que" und drückt Subjektivität aus.')

    # Categories
    doc.add_heading('3.1 Willensausdruck', level=2)
    doc.add_paragraph('Je veux que tu fasses tes devoirs.', style='List Bullet')
    doc.add_paragraph('J\'aimerais qu\'il vienne.', style='List Bullet')

    doc.add_heading('3.2 Notwendigkeit', level=2)
    doc.add_paragraph('Il faut que nous partions maintenant.', style='List Bullet')
    doc.add_paragraph('Il est important que vous appreniez le vocabulaire.', style='List Bullet')

    doc.add_heading('3.3 Gefühle & Emotionen', level=2)
    doc.add_paragraph('Je suis content que tu sois là.', style='List Bullet')
    doc.add_paragraph('Il a peur que le train ait du retard.', style='List Bullet')

    doc.add_heading('3.4 Konjunktionen (Auslöser)', level=2)
    doc.add_paragraph('Einige Konjunktionen verlangen zwingend den Subjonctif:', style='List Bullet')
    
    table_conj = doc.add_table(rows=1, cols=2)
    table_conj.style = 'Table Grid'
    table_conj.rows[0].cells[0].text = 'Konjunktion'
    table_conj.rows[0].cells[1].text = 'Beispiel'
    
    conj_data = [
        ('bien que (obwohl)', 'Bien qu\'il soit malade, il travaille.'),
        ('pour que (damit)', 'Je t\'aide pour que tu comprennes.'),
        ('sans que (ohne dass)', 'Il part sans que je le sache.')
    ]
    for c, ex in conj_data:
        r = table_conj.add_row().cells
        r[0].text = c
        r[1].text = ex

    # 4. Contrast
    doc.add_heading('4. Indicatif oder Subjonctif?', level=1)
    doc.add_paragraph('Verben des Denkens/Meinens (penser, croire, trouver) verhalten sich speziell:')
    
    p_ind = doc.add_paragraph(style='List Bullet')
    p_ind.add_run('Positiv = Indicatif: ').bold = True
    p_ind.add_run('Je pense qu\'il est intelligent.')
    
    p_sub = doc.add_paragraph(style='List Bullet')
    p_sub.add_run('Negativ/Frage = Subjonctif: ').bold = True
    p_sub.add_run('Je ne pense pas qu\'il soit intelligent.')

    # Save
    file_path = 'Lerninhalt_Franzoesisch_Subjonctif_Final.docx'
    doc.save(file_path)
    return file_path

create_final_subjonctif_docx()